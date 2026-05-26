"""从 .docx 提取纯文本。"""

from __future__ import annotations

import io
import os
from pathlib import Path

from docx import Document

ALLOWED_EXTENSIONS = {".docx"}

SECTION_QUESTION = "【题目卷】"
SECTION_ANSWER = "【答案解析卷】"


def combine_volume_texts(question_text: str, answer_text: str = "") -> str:
    parts = [f"{SECTION_QUESTION}\n{str(question_text or '').strip()}"]
    ans = str(answer_text or "").strip()
    if ans:
        parts.append(f"{SECTION_ANSWER}\n{ans}")
    return "\n\n".join(parts).strip()


def min_extract_chars() -> int:
    raw = os.getenv("QUESTION_IMPORT_MIN_EXTRACT_CHARS", "50").strip()
    try:
        return max(1, int(raw))
    except ValueError:
        return 50


def normalize_extension(filename: str) -> str:
    return Path(str(filename or "").strip()).suffix.lower()


def extract_docx_bytes(data: bytes) -> str:
    if not data:
        raise ValueError("文件为空")
    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise ValueError(f"无法读取 docx：{exc}") from exc

    parts: list[str] = []
    for para in doc.paragraphs:
        t = str(para.text or "").strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells: list[str] = []
            for cell in row.cells:
                cell_lines = [ln.strip() for ln in str(cell.text or "").splitlines() if ln.strip()]
                if cell_lines:
                    cells.append("\n".join(cell_lines))
            if cells:
                parts.append("\t".join(cells))

    return "\n\n".join(parts).strip()


def extract_document_bytes(data: bytes, *, ext: str) -> str:
    normalized = str(ext or "").strip().lower()
    if normalized == ".docx":
        return extract_docx_bytes(data)
    raise ValueError("仅支持 .docx 文件")


def validate_extracted_text(text: str, *, ext: str = "") -> None:
    del ext
    n = len(str(text or "").strip())
    if n < min_extract_chars():
        raise ValueError(
            f"文档无可识别文字（仅提取到 {n} 字，需要至少 {min_extract_chars()} 字）。"
            "请确认 docx 内为可复制文本。"
        )
